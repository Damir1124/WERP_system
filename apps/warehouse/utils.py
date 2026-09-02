"""
Утилиты модуля Warehouse.

- generate_waybill(courier_id, date) — генерация путевого листа .docx
  на основе модели Garage (автомобиль курьера) и CourierTrip (рейсы за дату).
"""
import io
import logging
from datetime import date

from django.utils import timezone

logger = logging.getLogger(__name__)


def generate_waybill(courier_id: int, date_str: str) -> str:
    """Генерация путевого листа .docx для курьера за дату.

    Возвращает путь к сохранённому файлу (в media/waybills/).
    """
    from docx import Document

    from apps.logistics.models import CourierShift, CourierTrip
    from apps.workers.models import Worker

    # Парсим дату
    if isinstance(date_str, str):
        target_date = date.fromisoformat(date_str)
    else:
        target_date = date_str

    # Курьер и его автомобиль
    courier = Worker.objects.get(pk=courier_id)
    garage = courier.garage if hasattr(courier, 'garage') else None

    # Смены курьера за дату и их рейсы
    shifts = CourierShift.objects.filter(courier=courier, date=target_date)
    trips = CourierTrip.objects.filter(shift__in=shifts)

    # Создаём документ
    doc = Document()
    doc.add_heading('Путевой лист', level=0)

    # Шапка
    doc.add_paragraph(f'Дата: {target_date.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Курьер: {courier.full_name}')
    if garage:
        doc.add_paragraph(f'Автомобиль: {garage.vehicle_name}')
        doc.add_paragraph(f'Гос. номер: {garage.plate_number or "—"}')
    doc.add_paragraph('')

    # Рейсы
    doc.add_heading('Рейсы', level=1)
    if not trips.exists():
        doc.add_paragraph('Рейсов за указанную дату нет.')
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Рейс'
        hdr[1].text = 'Загружено'
        hdr[2].text = 'Доставлено'
        hdr[3].text = 'Возвращено'
        hdr[4].text = 'Статус'

        for trip in trips:
            summary = trip.get_trip_summary()
            row = table.add_row().cells
            row[0].text = f'#{trip.id}'
            row[1].text = str(summary['full_loaded'])
            row[2].text = str(summary['delivered'])
            row[3].text = str(summary['full_returned'])
            row[4].text = trip.get_status_display()

    # Сохраняем в media/waybills/
    from django.conf import settings
    import os

    waybills_dir = os.path.join(settings.MEDIA_ROOT, 'waybills')
    os.makedirs(waybills_dir, exist_ok=True)

    filename = f'waybill_{courier_id}_{target_date.isoformat()}.docx'
    file_path = os.path.join(waybills_dir, filename)

    doc.save(file_path)
    logger.info("Путевой лист сохранён: %s", file_path)
    return file_path